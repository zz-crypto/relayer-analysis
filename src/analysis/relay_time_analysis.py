import mysql.connector
import matplotlib.pyplot as plt
from datetime import datetime
from docx import Document
from docx.shared import Inches
import io
import numpy as np
from collections import defaultdict
import pandas as pd
import json

with open("database_config.json") as f:
    db_config = json.load(f)


query = """
SELECT 
    rar.origin_chain_id,
    rar.destination_chain_id,
    rar.input_symbol,
    rar.output_symbol,
    CASE 
        WHEN rar.input_amount_usd < 1000 THEN '0-1k'
        WHEN rar.input_amount_usd < 10000 THEN '1k-10k'
        WHEN rar.input_amount_usd < 100000 THEN '10k-100k'
        ELSE '100k+'
    END AS amount_range,
    rar.relayer,
    rar.relay_time,
    rar.fill_block_time
FROM 
    relay_analysis_results rar
JOIN 
    target_relayer_combo trc
    ON rar.origin_chain_id = trc.origin_chain_id
    AND rar.destination_chain_id = trc.destination_chain_id
    AND rar.input_symbol = trc.input_symbol
    AND rar.output_symbol = trc.output_symbol
    AND rar.relayer = trc.relayer
    AND CASE 
        WHEN rar.input_amount_usd < 1000 THEN '0-1k'
        WHEN rar.input_amount_usd < 10000 THEN '1k-10k'
        WHEN rar.input_amount_usd < 100000 THEN '10k-100k'
        ELSE '100k+'
    END = trc.amount_range
"""


conn = mysql.connector.connect(**db_config)
cursor = conn.cursor()
cursor.execute(query)
results = cursor.fetchall()


doc = Document()


general_combos = defaultdict(lambda: defaultdict(list))
for row in results:
    general_combo = (row[0], row[1], row[2], row[3], row[4])
    specific_combo = (row[0], row[1], row[2], row[3], row[4], row[5])
    general_combos[general_combo][specific_combo].append(
        (datetime.strptime(str(row[7]), "%Y-%m-%d %H:%M:%S"), row[6])
    )


def remove_outliers(x, y):
    q1 = np.percentile(y, 25)
    q3 = np.percentile(y, 75)
    iqr = q3 - q1
    lower_bound = q1 - (1.5 * iqr)
    upper_bound = q3 + (1.5 * iqr)
    mask = (y >= lower_bound) & (y <= upper_bound)
    return np.array(x)[mask], np.array(y)[mask]


for general_combo, specific_combos in general_combos.items():
    doc.add_heading(f"General Combo: {general_combo}", level=1)

    plt.figure(figsize=(12, 6))

    all_data = defaultdict(list)
    for specific_combo, data in specific_combos.items():
        relayer = specific_combo[-1]
        for date, relay_time in data:
            all_data[relayer].append((date.date(), relay_time))

    for relayer, data in all_data.items():
        df = pd.DataFrame(data, columns=["date", "relay_time"])
        df = df.groupby("date")["relay_time"].mean().reset_index()
        df = df.sort_values("date")
        plt.plot(df["date"], df["relay_time"], label=relayer)

    plt.xlabel("Date")
    plt.ylabel("Average Relay Time")
    plt.title(f"Average Daily Relay Time for All Relayers - {general_combo}")
    plt.legend()
    plt.xticks(rotation=45)
    plt.tight_layout()

    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format="png")
    img_buffer.seek(0)

    doc.add_picture(img_buffer, width=Inches(6))
    doc.add_paragraph(
        f"Average daily relay time for all relayers in general combo: {general_combo}"
    )

    plt.close()

    for specific_combo, data in specific_combos.items():
        x = [d[0] for d in data]
        y = [d[1] for d in data]

        x_clean, y_clean = remove_outliers(x, y)

        plt.figure(figsize=(10, 6))
        plt.scatter(x_clean, y_clean)
        plt.xlabel("Fill Block Time")
        plt.ylabel("Relay Time")
        plt.title(f"Specific Combo: {specific_combo} (Outliers Removed)")
        plt.xticks(rotation=45)
        plt.tight_layout()

        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format="png")
        img_buffer.seek(0)

        doc.add_picture(img_buffer, width=Inches(6))
        doc.add_paragraph(
            f"Scatter plot for specific combo: {specific_combo} (Outliers Removed)"
        )

        plt.close()


doc.save("relay_time_analysis.docx")


cursor.close()
conn.close()

print("Relay time analysis has been saved to relay_time_analysis.docx")
