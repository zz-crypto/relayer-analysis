import mysql.connector
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from io import BytesIO
import json

with open("database_config.json") as f:
    db_config = json.load(f)

data_query = """
SELECT 
    r.origin_chain_id,
    r.destination_chain_id,
    r.input_symbol,
    r.output_symbol,
    CASE 
        WHEN r.output_amount_usd < 1000 THEN '0-1000'
        WHEN r.output_amount_usd < 10000 THEN '1000-10000'
        WHEN r.output_amount_usd < 100000 THEN '10000-100000'
        WHEN r.output_amount_usd < 1000000 THEN '100000-1000000'
        ELSE '1000000+'
    END AS output_amount_range,
    DATE(r.fill_block_time) AS date,
    CASE 
        WHEN HOUR(r.fill_block_time) < 12 THEN 'AM'
        ELSE 'PM'
    END AS half_day,
    AVG(r.priority_fee * 2700 / (r.input_amount_usd - r.output_amount_usd)) AS avg_priority_fee_ratio
FROM relay_analysis_results r
JOIN target_combo t ON 
    r.origin_chain_id = t.origin_chain_id AND
    r.destination_chain_id = t.destination_chain_id AND
    r.input_symbol = t.input_symbol AND
    r.output_symbol = t.output_symbol
WHERE r.priority_fee * 2700 / (r.input_amount_usd - r.output_amount_usd) < 1
GROUP BY 
    r.origin_chain_id,
    r.destination_chain_id,
    r.input_symbol,
    r.output_symbol,
    output_amount_range,
    DATE(r.fill_block_time),
    half_day
ORDER BY 
    r.origin_chain_id,
    r.destination_chain_id,
    r.input_symbol,
    r.output_symbol,
    output_amount_range,
    DATE(r.fill_block_time),
    half_day
"""
conn = mysql.connector.connect(**db_config)
cursor = conn.cursor()
cursor.execute(data_query)
data = cursor.fetchall()
columns = [
    "origin_chain_id",
    "destination_chain_id",
    "input_symbol",
    "output_symbol",
    "output_amount_range",
    "date",
    "half_day",
    "avg_priority_fee_ratio",
]
df = pd.DataFrame(data, columns=columns)
df["datetime"] = pd.to_datetime(df["date"]) + pd.to_timedelta(
    df["half_day"].map({"AM": "00:00:00", "PM": "12:00:00"})
)
doc = Document()
for (origin, dest, input_sym, output_sym), group in df.groupby(
    ["origin_chain_id", "destination_chain_id", "input_symbol", "output_symbol"]
):
    plt.figure(figsize=(12, 7))

    for amount_range, subgroup in group.groupby("output_amount_range"):
        plt.plot(
            subgroup["datetime"], subgroup["avg_priority_fee_ratio"], label=amount_range
        )

    plt.title(f"{origin} to {dest}: {input_sym} -> {output_sym}")
    plt.xlabel("Date")
    plt.ylabel("Average Priority Fee Ratio")
    plt.xticks(rotation=45)
    plt.legend()
    plt.tight_layout()

    img_buffer = BytesIO()
    plt.savefig(img_buffer, format="png")
    doc.add_picture(img_buffer, width=Inches(6))
    plt.close()

    doc.add_paragraph(f"{origin} to {dest}: {input_sym} -> {output_sym}")
doc.save("priority_fee_ratio_charts.docx")
conn.close()
