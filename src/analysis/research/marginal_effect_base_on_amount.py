import mysql.connector
import matplotlib.pyplot as plt
from matplotlib.ticker import FuncFormatter
from docx import Document
from docx.shared import Inches
import io
import numpy as np
import json

with open("database_config.json") as f:
    config = json.load(f)


query = """
SELECT 
    r.input_symbol,
    r.output_symbol,
    r.origin_chain_id,
    r.destination_chain_id,
    r.output_amount_usd * 4 AS amount,
    (input_amount_usd - output_amount_usd - gas_fee * 2700) / output_amount_usd / 4 AS profit_rate,
    (input_amount_usd - output_amount_usd - gas_fee * 2700) AS profit
FROM 
    relay_analysis_results r
INNER JOIN (
    SELECT DISTINCT
        origin_chain_id,
        destination_chain_id,
        input_symbol,
        output_symbol
    FROM 
        target_combo
) tc ON 
    r.origin_chain_id = tc.origin_chain_id AND
    r.destination_chain_id = tc.destination_chain_id AND
    r.input_symbol = tc.input_symbol AND
    r.output_symbol = tc.output_symbol
WHERE
    input_amount_usd - output_amount_usd - gas_fee * 2700 > 0 and output_amount_usd > 10000
ORDER BY 
    r.input_symbol,r. output_symbol, r.origin_chain_id, r.destination_chain_id, output_amount_usd;
"""

cnx = mysql.connector.connect(**config)
cursor = cnx.cursor()
cursor.execute(query)


results = cursor.fetchall()


cursor.close()
cnx.close()


data = {}
for row in results:
    key = (row[0], row[1], row[2], row[3])
    if key not in data:
        data[key] = {"amount": [], "profit": [], "profit_rate": []}
    data[key]["amount"].append(row[4])
    data[key]["profit"].append(row[6])
    data[key]["profit_rate"].append(row[5])


doc = Document()
doc.add_heading("Cumulative Profit and Average Profit Rate vs Amount Charts", 0)


for key, values in data.items():
    fig, ax1 = plt.subplots(figsize=(14, 7))

    sorted_indices = np.argsort(values["amount"])
    sorted_amounts = np.array(values["amount"])[sorted_indices]
    sorted_profits = np.array(values["profit"])[sorted_indices]

    cumulative_profits = np.cumsum(sorted_profits)

    average_rates = cumulative_profits / sorted_amounts

    color = "tab:blue"
    ax1.set_xlabel("Amount")
    ax1.set_ylabel("Cumulative Profit", color=color)
    ax1.plot(sorted_amounts, cumulative_profits, color=color)
    ax1.tick_params(axis="y", labelcolor=color)

    ax2 = ax1.twinx()
    color = "tab:orange"
    ax2.set_ylabel("Average Profit Rate", color=color)
    ax2.plot(sorted_amounts, average_rates, color=color)
    ax2.tick_params(axis="y", labelcolor=color)

    ax1.set_xscale("log")
    min_amount = sorted_amounts.min()
    max_amount = sorted_amounts.max()
    num_ticks = 15

    log_ticks = np.linspace(np.log10(min_amount), np.log10(max_amount), num_ticks)
    ticks = 10**log_ticks

    ax1.set_xticks(ticks)
    ax1.xaxis.set_major_formatter(FuncFormatter(lambda x, p: format(int(x), ",")))

    plt.xticks(rotation=45, ha="right")

    label = f"{key[0]}-{key[1]} ({key[2]}->{key[3]})"
    plt.title(f"Cumulative Profit and Average Profit Rate vs Amount for {label}")
    plt.grid(True)

    lines1, labels1 = ax1.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(
        lines1 + lines2, ["Cumulative Profit", "Avg Profit Rate"], loc="upper left"
    )

    plt.tight_layout()

    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format="png", dpi=300, bbox_inches="tight")
    img_buffer.seek(0)

    doc.add_picture(img_buffer, width=Inches(7))
    doc.add_paragraph(f"Chart for {label}")

    plt.close()


doc.save("cumulative_profit_and_rate_vs_amount_charts.docx")

print("All save to 'cumulative_profit_and_rate_vs_amount_charts.docx'")
