import mysql.connector
import pandas as pd
import matplotlib.pyplot as plt
from datetime import timedelta
import numpy as np
from decimal import Decimal
from docx import Document
from docx.shared import Inches
from io import BytesIO
import csv
import json


with open("database_config.json") as f:
    db_config = json.load(f)


def calculate_max_fund(df):
    df["time_window"] = df["fill_block_time"].dt.floor("2h")
    window_sums = df.groupby("time_window")["output_amount_usd"].sum()
    return Decimal(str(np.nanmax(window_sums.values + np.roll(window_sums.values, 1))))


def simulate(initial_fund, df):
    available_fund = initial_fund
    total_profit = Decimal("0")
    missed_orders = 0
    missed_profit = Decimal("0")
    last_refill_time = df["fill_block_time"].iloc[0]

    for _, row in df.iterrows():
        if row["fill_block_time"] - last_refill_time >= timedelta(hours=2):
            available_fund = initial_fund
            last_refill_time = row["fill_block_time"]

        if available_fund >= row["output_amount_usd"]:
            available_fund -= row["output_amount_usd"]
            total_profit += row["profit"]
        else:
            missed_orders += 1
            missed_profit += row["profit"]

    return total_profit, missed_orders, missed_profit


trade_pairs = [
    (1, 8453, "USDC", "USDC", 10000, 100000),
    (1, 42161, "USDC", "USDC", 10000, 100000),
    (1, 42161, "USDT", "USDT", 10000, 100000),
    (1, 42161, "WETH", "WETH", 10000, 100000),
    (8453, 1, "USDC", "USDC", 10000, 100000),
    (8453, 1, "WETH", "WETH", 10000, 100000),
    (42161, 1, "DAI", "DAI", 10000, 100000),
    (42161, 1, "USDC", "USDC", 10000, 100000),
    (42161, 1, "WBTC", "WBTC", 10000, 100000),
    (42161, 1, "WETH", "WETH", 10000, 100000),
]


doc = Document()
doc.add_heading("Trade Pair Analysis", 0)

csv_filename = "simulation_results.csv"
with open(csv_filename, "w", newline="") as csvfile:
    csvwriter = csv.writer(csvfile)
    csvwriter.writerow(
        [
            "origin_chain",
            "destination_chain",
            "input_symbol",
            "output_symbol",
            "min_amount",
            "max_amount",
            "simulated_allocation",
            "profit",
            "profit_ratio",
        ]
    )


conn = mysql.connector.connect(**db_config)
cursor = conn.cursor()

for (
    origin_chain,
    dest_chain,
    input_sym,
    output_sym,
    min_amount,
    max_amount,
) in trade_pairs:

    query = f"""
    SELECT output_amount_usd, fill_block_time, input_amount_usd, gas_fee
    FROM relay_analysis_results
    WHERE output_amount_usd BETWEEN {min_amount} AND {max_amount}
    AND origin_chain_id = '{origin_chain}' 
    AND destination_chain_id = '{dest_chain}' 
    AND input_symbol = '{input_sym}' 
    AND output_symbol = '{output_sym}'
    ORDER BY fill_block_time
    """

    cursor.execute(query)
    data = cursor.fetchall()

    df = pd.DataFrame(
        data,
        columns=["output_amount_usd", "fill_block_time", "input_amount_usd", "gas_fee"],
    )

    df["profit"] = (
        df["input_amount_usd"] - df["output_amount_usd"] - df["gas_fee"] * 2700
    )

    df["fill_block_time"] = pd.to_datetime(df["fill_block_time"])

    max_fund = calculate_max_fund(df)
    fund_levels = [max_fund * Decimal(str(1 - i * 0.05)) for i in range(20)]
    results = []

    for fund in fund_levels:
        total_profit, missed_orders, missed_profit = simulate(fund, df)
        profit_ratio = total_profit / fund if fund > 0 else Decimal("0")
        results.append(
            {
                "fund": fund,
                "total_profit": total_profit,
                "missed_orders": missed_orders,
                "missed_profit": missed_profit,
                "profit_loss_percentage": (
                    missed_profit / (total_profit + missed_profit) * Decimal("100")
                    if (total_profit + missed_profit) > 0
                    else Decimal("0")
                ),
                "profit_to_allocation_ratio": profit_ratio,
            }
        )

        with open(csv_filename, "a", newline="") as csvfile:
            csvwriter = csv.writer(csvfile)
            csvwriter.writerow(
                [
                    origin_chain,
                    dest_chain,
                    input_sym,
                    output_sym,
                    min_amount,
                    max_amount,
                    float(fund),
                    float(total_profit),
                    float(profit_ratio),
                ]
            )

    results_df = pd.DataFrame(results)
    results_df = results_df.applymap(
        lambda x: float(x) if isinstance(x, Decimal) else x
    )

    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(10, 12))

    ax1.plot(results_df["fund"], results_df["profit_loss_percentage"])
    ax1.set_xlabel("Initial Fund")
    ax1.set_ylabel("Profit Loss Percentage")
    ax1.set_title(
        f"Impact of Initial Fund on Profit Loss\n{origin_chain} to {dest_chain} - {input_sym} to {output_sym}"
    )
    ax1.grid(True)

    ax2.plot(results_df["fund"], results_df["profit_to_allocation_ratio"])
    ax2.set_xlabel("Initial Fund")
    ax2.set_ylabel("Profit to Allocation Ratio")
    ax2.set_title(
        f"Profit to Allocation Ratio vs Initial Fund\n{origin_chain} to {dest_chain} - {input_sym} to {output_sym}"
    )
    ax2.grid(True)

    plt.tight_layout()

    img_buffer = BytesIO()
    plt.savefig(img_buffer, format="png")
    img_buffer.seek(0)
    doc.add_heading(
        f"{origin_chain} to {dest_chain} - {input_sym} to {output_sym}", level=1
    )
    doc.add_picture(img_buffer, width=Inches(6))
    plt.close()

    best_ratio_point = results_df.loc[results_df["profit_to_allocation_ratio"].idxmax()]
    doc.add_paragraph(f"Best profit to allocation ratio point:")
    doc.add_paragraph(f"Initial fund: ${best_ratio_point['fund']:.2f}")
    doc.add_paragraph(
        f"Profit to allocation ratio: {best_ratio_point['profit_to_allocation_ratio']:.4f}"
    )
    doc.add_paragraph(f"Profit loss: {best_ratio_point['profit_loss_percentage']:.2f}%")

    doc.add_paragraph(
        "Figure 1 shows the impact of initial fund on profit loss. As the initial fund decreases, the percentage of profit loss increases."
    )
    doc.add_paragraph(
        "Figure 2 shows the profit to allocation ratio versus the initial fund. The peak of this curve represents the most efficient use of capital."
    )


cursor.close()
conn.close()


doc.save("trade_pair_analysis.docx")

print(
    f"Analysis complete. Results saved in 'trade_pair_analysis.docx' and '{csv_filename}'."
)
