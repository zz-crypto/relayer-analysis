import mysql.connector
import pandas as pd
import matplotlib.pyplot as plt
from datetime import timedelta
import numpy as np
from decimal import Decimal
from docx import Document
from docx.shared import Inches
from io import BytesIO
import json

with open("database_config.json") as f:
    db_config = json.load(f)


def calculate_max_fund(df):
    df["time_window"] = df["fill_block_time"].dt.floor("2h")
    window_sums = df.groupby("time_window")["output_amount_usd"].sum()
    return Decimal(str(np.nanmax(window_sums.values + np.roll(window_sums.values, 1))))


def simulate(initial_fund, df):
    available_fund = initial_fund
    total_profit = Decimal("0")
    missed_orders = 0
    missed_profit = Decimal("0")
    last_refill_time = df["fill_block_time"].iloc[0]

    for _, row in df.iterrows():
        if row["fill_block_time"] - last_refill_time >= timedelta(hours=2):
            available_fund = initial_fund
            last_refill_time = row["fill_block_time"]

        if available_fund >= row["output_amount_usd"]:
            available_fund -= row["output_amount_usd"]
            total_profit += row["profit"]
        else:
            missed_orders += 1
            missed_profit += row["profit"]

    return total_profit, missed_orders, missed_profit


trade_pairs = [
    (1, 8453, "USDC", "USDC"),
    (1, 42161, "USDC", "USDC"),
    (1, 42161, "USDT", "USDT"),
    (1, 42161, "WETH", "WETH"),
    (8453, 1, "USDC", "USDC"),
    (8453, 1, "WETH", "WETH"),
    (42161, 1, "DAI", "DAI"),
    (42161, 1, "USDC", "USDC"),
    (42161, 1, "WBTC", "WBTC"),
    (42161, 1, "WETH", "WETH"),
]


doc = Document()
doc.add_heading("Relayer Analysis", 0)


conn = mysql.connector.connect(**db_config)
cursor = conn.cursor()

for origin_chain, dest_chain, input_sym, output_sym in trade_pairs:

    query = f"""
    SELECT output_amount_usd, fill_block_time, input_amount_usd, gas_fee
    FROM relay_analysis_results
    WHERE origin_chain_id = '{origin_chain}' 
    AND destination_chain_id = '{dest_chain}' 
    AND input_symbol = '{input_sym}' 
    AND output_symbol = '{output_sym}'
    AND relayer = '0x07aE8551Be970cB1cCa11Dd7a11F47Ae82e70E67'
    ORDER BY fill_block_time
    """

    cursor.execute(query)
    data = cursor.fetchall()

    if not data:
        continue

    df = pd.DataFrame(
        data,
        columns=["output_amount_usd", "fill_block_time", "input_amount_usd", "gas_fee"],
    )

    df["profit"] = (
        df["input_amount_usd"] - df["output_amount_usd"] - df["gas_fee"] * 2700
    )

    df["fill_block_time"] = pd.to_datetime(df["fill_block_time"])

    max_fund = calculate_max_fund(df)
    fund_levels = [max_fund * Decimal(str(1 - i * 0.05)) for i in range(20)]
    results = []

    for fund in fund_levels:
        total_profit, missed_orders, missed_profit = simulate(fund, df)
        results.append(
            {
                "fund": fund,
                "total_profit": total_profit,
                "missed_orders": missed_orders,
                "missed_profit": missed_profit,
                "profit_to_allocation_ratio": (
                    total_profit / fund if fund > 0 else Decimal("0")
                ),
            }
        )

    results_df = pd.DataFrame(results)
    results_df = results_df.applymap(
        lambda x: float(x) if isinstance(x, Decimal) else x
    )

    plt.figure(figsize=(10, 6))
    plt.plot(results_df["fund"], results_df["profit_to_allocation_ratio"])
    plt.xlabel("Initial Fund")
    plt.ylabel("Profit to Allocation Ratio")
    plt.title(
        f"Profit to Allocation Ratio vs Initial Fund\n{origin_chain} to {dest_chain} - {input_sym} to {output_sym}"
    )
    plt.grid(True)

    img_buffer = BytesIO()
    plt.savefig(img_buffer, format="png")
    img_buffer.seek(0)
    doc.add_heading(
        f"{origin_chain} to {dest_chain} - {input_sym} to {output_sym}", level=1
    )
    doc.add_picture(img_buffer, width=Inches(6))
    plt.close()

    best_ratio_point = results_df.loc[results_df["profit_to_allocation_ratio"].idxmax()]
    doc.add_paragraph(f"Best profit to allocation ratio point:")
    doc.add_paragraph(f"Initial fund: ${best_ratio_point['fund']:.2f}")
    doc.add_paragraph(
        f"Profit to allocation ratio: {best_ratio_point['profit_to_allocation_ratio']:.4f}"
    )

    max_allocation_profit_ratio = results_df.iloc[0]["profit_to_allocation_ratio"]
    doc.add_paragraph(
        f"Profit to allocation ratio at max allocation: {max_allocation_profit_ratio:.4f}"
    )

    max_allocation = results_df.iloc[0]["fund"]
    doc.add_paragraph(f"Max allocation: ${max_allocation:.2f}")

    doc.add_paragraph(
        "The graph shows the profit to allocation ratio versus the initial fund. The peak of this curve represents the most efficient use of capital."
    )

cursor.close()
conn.close()

doc.save("relayer_analysis.docx")
