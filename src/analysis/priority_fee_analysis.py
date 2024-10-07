
import mysql.connector
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import io
import numpy as np
import json

with open("database_config.json") as f:
    db_config = json.load(f)

conn = mysql.connector.connect(**db_config)
cursor = conn.cursor()

query = """
SELECT 
    r.origin_chain_id,
    r.destination_chain_id,
    r.input_symbol,
    r.output_symbol,
    r.gas_fee * 2700 AS gas_fee_in_usd,
    r.priority_fee * 2700 / (r.input_amount_usd - r.output_amount_usd) AS priority_fee_ratio,
    r.input_amount_usd - r.output_amount_usd - r.gas_fee * 2700 AS net_profit,
    r.input_amount_usd,
    r.output_amount_usd
FROM relay_analysis_results r
JOIN target_combo t ON 
    r.origin_chain_id = t.origin_chain_id AND
    r.destination_chain_id = t.destination_chain_id AND
    r.input_symbol = t.input_symbol AND
    r.output_symbol = t.output_symbol
WHERE r.priority_fee * 2700 / (r.input_amount_usd - r.output_amount_usd) < 1
"""

df = pd.read_sql(query, conn)
def find_optimal_percentile(data):
    best_percentile = 0
    max_expected_profit = float('-inf')
    for p in range(1, 100):
        ratio = data['priority_fee_ratio'].quantile(p/100)
        success_rate = 1 - p/100
        
        
        success_profit = (data['net_profit'] - (ratio * (data['input_amount_usd'] - data['output_amount_usd']))) * success_rate
        
        
        failure_loss = -data['gas_fee_in_usd'] * (1 - success_rate)
        
        expected_profit = success_profit + failure_loss
        avg_expected_profit = expected_profit.mean()
        
        if avg_expected_profit > max_expected_profit:
            max_expected_profit = avg_expected_profit
            best_percentile = p
            best_ratio = ratio
    
    return best_percentile, max_expected_profit, best_ratio

doc = Document()
doc.add_heading('Optimal Priority Fee Analysis Report', 0)

combos = df.groupby(['origin_chain_id', 'destination_chain_id', 'input_symbol', 'output_symbol'])

log_ranges = [(10**i, 10**(i+1)) for i in range(0, 6)]  

for name, group in combos:
    origin_chain_id, destination_chain_id, input_symbol, output_symbol = name
    
    doc.add_paragraph(f'Combo: {input_symbol}->{output_symbol} ({origin_chain_id}->{destination_chain_id})')
    
    plt.figure(figsize=(10, 6))
    plt.scatter(group['input_amount_usd'], group['priority_fee_ratio'], alpha=0.5, s=10)
    plt.title(f'Optimal Priority Fee Ratio vs Input Amount\n{input_symbol}->{output_symbol} ({origin_chain_id}->{destination_chain_id})')
    plt.xlabel('Input Amount (USD)')
    plt.ylabel('Priority Fee Ratio')
    plt.xscale('log')
    plt.ylim(0, min(1, group['priority_fee_ratio'].quantile(0.99)))
    
    for start, end in log_ranges:
        range_data = group[(group['input_amount_usd'] >= start) & (group['input_amount_usd'] < end)]
        if not range_data.empty:
            best_percentile, max_expected_profit, optimal_ratio = find_optimal_percentile(range_data)
            
            plt.hlines(optimal_ratio, start, end, colors='r', linestyles='dashed')
            plt.text(np.sqrt(start*end), optimal_ratio, f'{optimal_ratio:.4f}', 
                     verticalalignment='bottom', horizontalalignment='center', fontsize=8)
    
    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
    img_buffer.seek(0)
    doc.add_picture(img_buffer, width=Inches(6))
    plt.close()
    
    doc.add_paragraph(f'Number of transactions: {len(group)}')
    
    for start, end in log_ranges:
        range_data = group[(group['input_amount_usd'] >= start) & (group['input_amount_usd'] < end)]
        if not range_data.empty:
            best_percentile, max_expected_profit, optimal_ratio = find_optimal_percentile(range_data)
            
            doc.add_paragraph(f'Range ${start:,.0f} - ${end:,.0f}:')
            doc.add_paragraph(f'  Optimal percentile: {best_percentile}')
            doc.add_paragraph(f'  Optimal priority fee ratio: {optimal_ratio:.4f}')
            doc.add_paragraph(f'  Expected average profit: ${max_expected_profit:.2f}')
            doc.add_paragraph(f'  Number of transactions: {len(range_data)}')
    
    doc.add_page_break()

doc.save('optimal_priority_fee_analysis.docx')

cursor.close()
conn.close()

print("Report generated: optimal_priority_fee_analysis.docx")